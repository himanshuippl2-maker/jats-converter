#!/usr/bin/env python3
"""DOCX → JATS XML v4 — PMC Compliant
Features:
  - Tables inline where first mentioned in body
  - Table xref: handles "Table N" split across bold runs
  - CrossRef API (by DOI, then by query)
  - PubMed/NLM eUtils API for enrichment
  - Proper reference parsing: authors, title, journal, pages
"""
import sys, re, json, hashlib, urllib.request, urllib.parse, time
from docx import Document

# ── IDs ──────────────────────────────────────────────────────
_ctr = [0]
def make_prefix(title=""):
    h = hashlib.md5((title or "article").encode()).hexdigest()[:8]
    # XML IDs must start with letter
    return 'id' + h
def nid(pfx, label):
    _ctr[0] += 1
    suf = hashlib.md5(f"{label}-{_ctr[0]}".encode()).hexdigest()[:12]
    return f"{label}-{pfx}-{_ctr[0]}-{suf}"
def xe(t):
    if not t: return ''
    return str(t).replace('&','&amp;').replace('<','&lt;').replace('>','&gt;').replace('"','&quot;')
def slugify(t):
    return re.sub(r'[^a-z0-9]+','-',t.lower()).strip('-')

# ── SEC-TYPE ─────────────────────────────────────────────────
SEC_MAP = {
    'introduction':'intro','material and methods':'methods',
    'materials and methods':'methods','methods':'methods','methodology':'methods',
    'results':'results','discussion':'discussion',
    'conclusion':'conclusions','conclusions':'conclusions',
    'acknowledgement':'acknowledgments','acknowledgements':'acknowledgments',
    'acknowledgment':'acknowledgments','acknowledgments':'acknowledgments',
}
def get_sec_type(title):
    lc = title.lower().strip()
    for k,v in SEC_MAP.items():
        if k in lc: return v
    return None

# ── AUTHOR PARSING ───────────────────────────────────────────
def parse_author_name(raw):
    name = raw.strip().rstrip('.,')
    if not name: return {'surname':'','given':'','prefix':''}
    parts = name.split()
    if len(parts) == 1: return {'surname':parts[0],'given':'','prefix':''}
    # "Surname AB" or "AB Surname"
    if re.match(r'^[A-Z]{1,3}\.?$', parts[-1]):
        return {'surname':' '.join(parts[:-1]),'given':parts[-1].rstrip('.'),'prefix':''}
    if re.match(r'^[A-Z]{1,3}\.?$', parts[0]):
        return {'surname':' '.join(parts[1:]),'given':parts[0].rstrip('.'),'prefix':''}
    return {'surname':parts[-1],'given':' '.join(parts[:-1]),'prefix':''}

def parse_authors_para(para):
    authors=[]; cur_name=''; cur_affs=[]; is_corr=False
    for run in para.runs:
        t = run.text
        if not t: continue
        if run.font.superscript:
            nums = [n.strip() for n in re.split(r'[,،\s]+',t) if n.strip().isdigit()]
            cur_affs.extend(nums)
        elif t.strip() == '*':
            is_corr = True
        elif ',' in t:
            parts = t.split(',')
            cur_name += parts[0]
            _flush(cur_name, cur_affs, is_corr, authors)
            cur_name = ','.join(parts[1:]).strip()
            cur_affs = []; is_corr = False
        else:
            cur_name += t
    if cur_name.strip(): _flush(cur_name, cur_affs, is_corr, authors)
    return authors

def _flush(name, affs, corr, out):
    name = name.strip().strip(',').strip()
    if not name: return
    pn = parse_author_name(name)
    if pn['surname']: out.append({**pn, 'affiliationNums':affs[:], 'isCorresponding':corr})

def parse_affiliation(para):
    num=''; parts=[]
    for run in para.runs:
        if run.font.superscript and run.text.strip().isdigit():
            num = run.text.strip()
        else:
            parts.append(run.text)
    full = ''.join(parts).strip()
    if not num:
        m = re.match(r'^(\d+)\s*(.+)', para.text.strip())
        if m: num, full = m.group(1), m.group(2).strip()
        else: full = para.text.strip()
    return num, full

# ── REFERENCE PARSING ────────────────────────────────────────
def parse_ref(raw):
    r = {'authors':[],'title':'','journal':'','year':'',
         'volume':'','issue':'','fpage':'','lpage':'',
         'doi':'','hasEtAl':False,'pubType':'journal'}

    # DOI
    for pat in [r'https?://doi\.org/(10\.[^\s,;\]]+)', r'\bDOI:?\s*(10\.[^\s,;\]]+)', r'\b(10\.\d{4,}/[^\s,;\]]+)']:
        dm = re.search(pat, raw, re.I)
        if dm:
            doi_val = dm.group(1).rstrip('.,]')
            # DOI must contain "/" to be valid (prefix/suffix)
            if '/' in doi_val:
                r['doi'] = doi_val
            break

    clean = re.sub(r'https?://\S+','',raw).strip()
    if re.search(r'\[dissertation\]|\[thesis\]', clean, re.I): r['pubType']='thesis'
    if re.search(r'\bbook\b|\bmonograph\b|\bpublisher\b', clean, re.I): r['pubType']='book'

    # Year
    ym = re.search(r'\b((?:19|20)\d{2})\b', clean)
    if ym: r['year'] = ym.group(1)

    # Pages vol(issue):fp-lp
    pm = re.search(r'(\d+)\s*\((\d+)\)\s*:\s*(\d+)\s*[-–]\s*(\d+)', clean)
    if pm:
        r['volume'],r['issue'],r['fpage'],r['lpage'] = pm.group(1),pm.group(2),pm.group(3),pm.group(4)
    else:
        pm2 = re.search(r'(\d+)\s*:\s*(\d+)\s*[-–]\s*(\d+)', clean)
        if pm2: r['volume'],r['fpage'],r['lpage'] = pm2.group(1),pm2.group(2),pm2.group(3)

    _fix_lpage(r)

    # Authors: "Surname AB, Surname CD, Surname EF."
    # Split on ". " before a capital that starts title
    m_split = re.match(
        r'^((?:[A-Z][a-zA-Z\-\']+\s+[A-Z]{1,3}(?:\s+[A-Z]{1,3})?(?:,\s*)?)+?(?:,?\s*(?:et al\.?|and\s+[A-Z]\w+\s+[A-Z]{1,3}))?)[\.,]\s+([A-Z].+)',
        clean)
    if m_split:
        author_str, rest_str = m_split.group(1), m_split.group(2)
        for ap in re.split(r',\s*', author_str):
            ap = ap.strip().rstrip('.')
            if re.search(r'et al\.?$', ap, re.I): r['hasEtAl']=True; break
            pn = parse_author_name(ap)
            if pn['surname'] and len(pn['surname']) > 1: r['authors'].append(pn)
        # Title: up to next ". Journal"
        tm = re.match(r'^(.+?)\.\s+([A-Z][^\d]{3,}(?:J\b|Journal|Ann|Rev|Int|Eur|Am|Br|Clin|Med|Sci|Res|Arch|Bull|Acad|Proc))', rest_str)
        if tm:
            r['title'] = tm.group(1).strip()
            r['journal'] = tm.group(2).split('.')[0].strip()
        else:
            r['title'] = rest_str.split('.')[0].strip()
    else:
        # Fallback: first sentence = authors+title
        r['title'] = clean.split('.')[0].strip()

    return r

def _fix_lpage(r):
    if r['fpage'] and r['lpage']:
        try:
            fp, lp = int(r['fpage']), int(r['lpage'])
            if lp < fp:
                prefix = str(fp)[:len(str(fp))-len(str(lp))]
                r['lpage'] = prefix + str(lp)
        except: pass

# ── CROSSREF API ─────────────────────────────────────────────
def fetch_crossref(doi=None, query=None, timeout=8):
    try:
        if doi:
            url = f"https://api.crossref.org/works/{urllib.parse.quote(doi)}"
        else:
            url = (f"https://api.crossref.org/works?query={urllib.parse.quote(query)}"
                   f"&rows=1&select=DOI,author,title,published,container-title,volume,issue,page")
        req = urllib.request.Request(url, headers={'User-Agent':'JATSConverter/4.0 (medical@example.com)'})
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            data = json.loads(resp.read())
        item = data['message'] if doi else (data.get('message',{}).get('items') or [None])[0]
        if not item: return None
        pages = item.get('page','')
        fp = lp = ''
        if '-' in pages:
            fp,lp = pages.split('-',1)
            try:
                fpi,lpi = int(fp.strip()),int(lp.strip())
                if lpi < fpi:
                    pre = str(fpi)[:len(str(fpi))-len(str(lpi))]
                    lp = pre+str(lpi)
            except: pass
        else: fp = pages
        return {
            'doi': item.get('DOI',''),
            'title': (item.get('title') or [''])[0],
            'journal': (item.get('container-title') or [''])[0],
            'year': str(((item.get('published',{}).get('date-parts') or [['']])[0] or [''])[0]),
            'volume': item.get('volume',''), 'issue': item.get('issue',''),
            'fpage': fp.strip(), 'lpage': lp.strip(),
            'authors': [{'surname':a.get('family',''),'given':a.get('given',''),'orcid':a.get('ORCID','').replace('http://orcid.org/','')}
                        for a in item.get('author',[])],
        }
    except: return None

# ── PUBMED/NLM API ───────────────────────────────────────────
def fetch_pubmed(doi=None, query=None, timeout=8):
    try:
        base = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils"
        q = f"{doi}[AID]" if doi else query
        surl = f"{base}/esearch.fcgi?db=pubmed&term={urllib.parse.quote(q)}&retmax=1&retmode=json"
        req = urllib.request.Request(surl, headers={'User-Agent':'JATSConverter/4.0'})
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            sdata = json.loads(resp.read())
        ids = sdata.get('esearchresult',{}).get('idlist',[])
        if not ids: return None
        pmid = ids[0]
        furl = f"{base}/esummary.fcgi?db=pubmed&id={pmid}&retmode=json"
        req2 = urllib.request.Request(furl, headers={'User-Agent':'JATSConverter/4.0'})
        with urllib.request.urlopen(req2, timeout=timeout) as resp2:
            fdata = json.loads(resp2.read())
        doc = fdata.get('result',{}).get(pmid,{})
        if not doc: return None
        authors = []
        for a in doc.get('authors',[]):
            nm = a.get('name','').strip()
            if ' ' in nm:
                sur,giv = nm.split(' ',1)
                authors.append({'surname':sur,'given':giv,'orcid':''})
            elif nm:
                authors.append({'surname':nm,'given':'','orcid':''})
        pages = doc.get('pages','')
        fp = lp = ''
        if '-' in pages:
            fp,lp = pages.split('-',1)
            try:
                fpi,lpi = int(fp),int(lp)
                if lpi < fpi:
                    pre = str(fpi)[:len(str(fpi))-len(str(lpi))]
                    lp = pre+str(lpi)
            except: pass
        else: fp = pages
        doi_out = ''
        for loc in doc.get('articleids',[]):
            if loc.get('idtype') == 'doi': doi_out = loc.get('value',''); break
        return {
            'pmid': pmid, 'doi': doi_out,
            'title': doc.get('title','').rstrip('.'),
            'journal': doc.get('fulljournalname','') or doc.get('source',''),
            'year': doc.get('pubdate','')[:4],
            'volume': doc.get('volume',''), 'issue': doc.get('issue',''),
            'fpage': fp.strip(), 'lpage': lp.strip(), 'authors': authors,
        }
    except: return None

# ── INLINE TEXT → JATS ───────────────────────────────────────
def para_to_inline(para, pfx, cite_ids, table_ids, fig_ids):
    """
    Smart inline conversion:
    1. Reconstruct full para text from runs
    2. Find Table/Fig refs in full text (may span multiple runs)
    3. Process run-by-run, emitting xref where a span starts
    """
    # Build run list
    runs = []
    for run in para.runs:
        if run.text:
            runs.append({'t': run.text, 'sup': bool(run.font.superscript),
                         'ital': bool(run.italic), 'bold': bool(run.bold)})

    # Full text for span detection
    full = ''.join(r['t'] for r in runs)

    # Find Table/Figure spans → map char_start → xref_xml
    span_starts = {}  # char_pos_start → (char_pos_end, xref_xml)
    for m in re.finditer(r'(Table|Figure|Fig\.?)\s+(\d+)', full):
        ref_t = 'table' if m.group(1) == 'Table' else 'fig'
        num = m.group(2)
        rid = table_ids.get(num) if ref_t == 'table' else fig_ids.get(num)
        if rid:
            xid = nid(pfx, 'xref')
            xref = f'<xref id="{xid}" rid="{rid}" ref-type="{ref_t}">{xe(m.group(0))}</xref>'
            span_starts[m.start()] = (m.end(), xref)

    out = ''; char_pos = 0; skip_to = -1

    for rd in runs:
        t = rd['t']
        run_start = char_pos
        run_end = char_pos + len(t)

        result = ''
        i = 0  # position within this run
        while i < len(t):
            abs_i = run_start + i

            # Are we inside a skip zone?
            if abs_i < skip_to:
                i += 1; continue

            # Does a span start here?
            if abs_i in span_starts:
                span_end, xref_xml = span_starts[abs_i]
                result += xref_xml
                skip_to = span_end
                i += 1; continue

            # Normal character — accumulate
            j = i + 1
            while j < len(t):
                abs_j = run_start + j
                if abs_j >= skip_to and abs_j in span_starts:
                    break
                if abs_j < skip_to:
                    j += 1; continue
                j += 1
            chunk = t[i:j]
            result += _fmt_chunk(chunk, rd, pfx)
            i = j

        out += result
        char_pos = run_end

    return out

def _fmt_chunk(t, rd, pfx):
    if not t or not t.strip(): return xe(t)
    sup = rd['sup']; ital = rd['ital']; bold = rd['bold']
    if sup:
        nums_str = t.strip()
        if re.match(r'^[\d,\s\-–]+$', nums_str):
            result = ''
            for num in re.split(r'[,\s]+', nums_str):
                num = num.strip()
                if num and num.isdigit():
                    xid = nid(pfx, 'x')
                    result += f'<xref id="{xid}" rid="{pfx}-B{num}" ref-type="bibr">{xe(num)}</xref>'
            return result if result else f'<sup>{xe(t)}</sup>'
        return f'<sup>{xe(t)}</sup>'
    elif ital and bold: return f'<bold><italic>{xe(t)}</italic></bold>'
    elif ital: return f'<italic>{xe(t)}</italic>'
    elif bold:
        bid = nid(pfx, 's')
        return f'<bold id="{bid}">{xe(t)}</bold>'
    else: return xe(t)

# ── DOCX PARSER ──────────────────────────────────────────────
def parse_docx(path, use_crossref=True):
    doc = Document(path)
    parsed = {
        'title':'','authors':[],'affiliations':{},'abstract':{},'keywords':[],'figures':[],
        'receivedDate':'','acceptedDate':'','sections':[],'references':[],'tables':[],
    }
    cur_sec = cur_sub = None; ref_num = 1; tbl_caption_num = 0

    for para in doc.paragraphs:
        sty = para.style.name; txt = para.text.strip()

        if sty == 'Title':
            parsed['title'] = txt

        elif sty == 'Author Name':
            parsed['authors'] = parse_authors_para(para)

        elif sty in ('Authors affiliation','Last Authors affiliation'):
            if txt:
                num, aff = parse_affiliation(para)
                if num: parsed['affiliations'][num] = aff
                else:
                    n = str(len(parsed['affiliations'])+1)
                    parsed['affiliations'][n] = txt

        elif sty == 'Abstract':
            if not txt or 'Open Access' in txt or txt.lower().startswith('for reprint'): continue
            lm = re.match(r'^(Background|Methods?|Results?|Conclusion|Objective|Aim|Discussion|Summary|Purpose):\s*', txt, re.I)
            if lm:
                parsed['abstract'][lm.group(1)] = txt[len(lm.group(0)):]
            else:
                k = list(parsed['abstract'].keys())[-1] if parsed['abstract'] else 'text'
                if k not in parsed['abstract']: parsed['abstract'][k] = txt
                else: parsed['abstract'][k] += ' ' + txt

        elif sty == 'Keywords':
            if 'Keywords:' in txt:
                parsed['keywords'] = [k.strip().rstrip('.') for k in txt.split('Keywords:',1)[1].split(',') if k.strip()]
            elif 'Received:' in txt:
                rm = re.search(r'Received:\s*([\d\-/]+)', txt)
                am = re.search(r'Accepted:\s*([\d\-/]+)', txt)
                if rm: parsed['receivedDate'] = rm.group(1)
                if am: parsed['acceptedDate'] = am.group(1)

        elif sty == 'Heading 1':
            lc = txt.lower()
            skip_kw = ['reference','source of fund','conflict','ethical','patient consent']
            if any(s in lc for s in skip_kw): cur_sec = cur_sub = None
            else:
                cur_sec = {'title':txt,'paragraphs':[],'subsections':[],'sec_type':get_sec_type(txt)}
                cur_sub = None; parsed['sections'].append(cur_sec)

        elif sty == 'Heading 2':
            if cur_sec:
                cur_sub = {'title':txt,'paragraphs':[],'sec_type':get_sec_type(txt)}
                cur_sec['subsections'].append(cur_sub)

        elif sty in ('Paragraph 1','2nd Para','List Paragraph','Normal'):
            if not txt: continue
            if cur_sub: cur_sub['paragraphs'].append(para)
            elif cur_sec: cur_sec['paragraphs'].append(para)

        elif sty == 'Reference':
            if txt and not txt.startswith('http'):
                rp = parse_ref(txt)
                parsed['references'].append({
                    'num': ref_num, 'raw': txt,
                    'doi': rp.get('doi',''), 'parsed': rp,
                    'crossref': None, 'pubmed': None
                })
                ref_num += 1
            elif txt.startswith('http') and parsed['references']:
                dm = re.search(r'10\.\S+', txt)
                if dm: parsed['references'][-1]['doi'] = dm.group(0).rstrip('.')

        elif sty == 'Table caption':
            if txt:
                tm = re.search(r'Table\s+(\d+)', txt, re.I)
                tnum = int(tm.group(1)) if tm else (tbl_caption_num+1)
                tbl_caption_num = tnum
                found = False
                for tbl in parsed['tables']:
                    if tbl['num'] == tnum: tbl['caption']=txt; found=True; break
                if not found:
                    parsed['tables'].append({'num':tnum,'caption':txt,'rows':[],'colwidths':[],'placed':False})

        elif sty in ('Caption', 'Figure Caption', 'Image Caption'):
            if txt:
                fm = re.search(r'(Fig\.?\s*|Figure\s*)(\d+)', txt, re.I)
                fnum = int(fm.group(2)) if fm else (len(parsed['figures'])+1)
                found = False
                for fig in parsed['figures']:
                    if fig['num'] == fnum: fig['caption']=txt; found=True; break
                if not found:
                    parsed['figures'].append({'num':fnum,'caption':txt,'placed':False,'has_image':False})

    # Detect inline images and associate with figure captions
    fig_counter = 0
    for i, para in enumerate(doc.paragraphs):
        xml_str = para._element.xml
        has_image = ('drawing' in xml_str or 'pic:' in xml_str or 'blipFill' in xml_str)
        if has_image:
            # Look for caption in adjacent paragraphs
            for offset in [1, 2, -1]:
                adj_idx = i + offset
                if 0 <= adj_idx < len(doc.paragraphs):
                    adj = doc.paragraphs[adj_idx]
                    adj_txt = adj.text.strip()
                    if re.search(r'(Fig\.?\s*|Figure\s*)\d+', adj_txt, re.I):
                        fm = re.search(r'(Fig\.?\s*|Figure\s*)(\d+)', adj_txt, re.I)
                        fnum = int(fm.group(2)) if fm else (fig_counter+1)
                        fig_counter += 1
                        # Check if already exists
                        found = any(f['num']==fnum for f in parsed['figures'])
                        if not found:
                            parsed['figures'].append({'num':fnum,'caption':adj_txt,'placed':False,'has_image':True})
                        else:
                            for f in parsed['figures']:
                                if f['num']==fnum: f['has_image']=True
                        break

    # Parse actual tables from docx
    tbl_num = 0
    for table in doc.tables:
        tbl_num += 1
        rows=[]; colwidths=[]
        try:
            total = sum(c.width or 1 for c in table.columns if c.width) or 1
            for col in table.columns:
                colwidths.append(round(((col.width or 1)/total)*100,2))
        except: pass
        for row in table.rows:
            cells=[]
            for cell in row.cells:
                tc = cell._tc
                ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                gs = tc.find(f'.//{{{ns}}}gridSpan')
                colspan = int(gs.get(f'{{{ns}}}val',1)) if gs is not None else 1
                cells.append({'text':cell.text.strip(),'colspan':colspan})
            rows.append(cells)
        found = False
        for tbl in parsed['tables']:
            if tbl['num'] == tbl_num and not tbl.get('rows'):
                tbl['rows']=rows; tbl['colwidths']=colwidths; found=True; break
        if not found:
            parsed['tables'].append({'num':tbl_num,'caption':'','rows':rows,'colwidths':colwidths,'placed':False})

    # API enrichment
    if use_crossref and parsed['references']:
        print(f"\n🔍 Enriching {len(parsed['references'])} references via CrossRef + PubMed...", file=sys.stderr)
        for ref in parsed['references']:
            p = ref['parsed']; doi = ref.get('doi','')
            cr = pm = None

            # 1) CrossRef by DOI
            if doi: cr = fetch_crossref(doi=doi)

            # 2) CrossRef by query if no DOI or failed
            if not cr:
                parts = []
                if p['authors']: parts.append(p['authors'][0].get('surname',''))
                if p['title']: parts.extend(p['title'].split()[:6])
                if p['year']: parts.append(p['year'])
                q = ' '.join(filter(None,parts))
                if q.strip(): cr = fetch_crossref(query=q)

            if cr and not doi: ref['doi'] = cr.get('doi','')

            # 3) PubMed
            doi_for_pm = ref.get('doi','')
            if doi_for_pm:
                pm = fetch_pubmed(doi=doi_for_pm)
            if not pm and p['authors'] and p['title']:
                q_pm = f"{p['authors'][0].get('surname','')} {p['title'][:40]} {p['year']}"
                pm = fetch_pubmed(query=q_pm)
            if pm:
                if not ref['doi'] and pm.get('doi'): ref['doi'] = pm['doi']

            ref['crossref'] = cr; ref['pubmed'] = pm
            sym = '✓CR' if cr else ('✓PM' if pm else '✗ ')
            print(f"  [{ref['num']:2d}] {sym} {ref['raw'][:60]}...", file=sys.stderr)
            time.sleep(0.12)

    return parsed

# ── XML BUILDER ──────────────────────────────────────────────
def build_xml(parsed, jm):
    pfx = make_prefix(parsed.get('title',''))
    _ctr[0] = 0

    cite_ids = {str(r['num']) for r in parsed.get('references',[])}
    table_ids = {str(t['num']): f"tw-{t['num']}" for t in parsed.get('tables',[])}
    fig_ids = {str(f['num']): f'fig-{f["num"]}' for f in parsed.get('figures',[])}

    L = []
    L += ['<?xml version="1.0" encoding="UTF-8"?>',
          '<!DOCTYPE article PUBLIC "-//NLM//DTD JATS (Z39.96) Journal Publishing DTD v1.2 20190208//EN"',
          '  "JATS-journalpublishing1-2.dtd">',
          '<article xmlns:xlink="http://www.w3.org/1999/xlink"',
          '         xmlns:ali="http://www.niso.org/schemas/ali/1.0/"',
          f'         article-type="{jm.get("articleType","research-article")}"',
          '         xml:lang="en">','','  <front>']

    # journal-meta
    jm_id = nid(pfx,'jm')
    L += [f'    <journal-meta id="{jm_id}">',
          f'      <journal-id journal-id-type="nlm-ta">{xe(jm.get("publisher","IP Innovative Publication"))}</journal-id>',
          f'      <journal-id journal-id-type="publisher-id">{xe(jm.get("publisher","IP Innovative Publication"))}</journal-id>',
          '      <journal-title-group>',
          f'        <journal-title>{xe(jm.get("name",""))}</journal-title>',
          '      </journal-title-group>']
    if jm.get('issnPrint'): L.append(f'      <issn publication-format="print">{xe(jm["issnPrint"].strip())}</issn>')
    if jm.get('issnElec'): L.append(f'      <issn publication-format="electronic">{xe(jm["issnElec"].strip())}</issn>')
    if jm.get('journalUrl'):
        L.append(f'      <self-uri xlink:href="{xe(jm["journalUrl"])}" xlink:title="{xe(jm.get("name","Journal"))}"/>')
    L += ['    </journal-meta>','']

    # article-meta
    am_id = nid(pfx,'am')
    L.append(f'    <article-meta id="{am_id}">')
    if jm.get('doi'): L.append(f'      <article-id pub-id-type="doi">{xe(jm["doi"])}</article-id>')
    L += ['      <article-categories>','        <subj-group subj-group-type="heading">',
          f'          <subject>{xe(jm.get("articleTypeLabel","Original Research Article"))}</subject>',
          '        </subj-group>','      </article-categories>','',
          '      <title-group>',f'        <article-title>{xe(parsed.get("title",""))}</article-title>',
          '      </title-group>','','      <contrib-group>']

    for auth in parsed.get('authors',[]):
        corr = ' corresp="yes"' if auth.get('isCorresponding') else ''
        L.append(f'        <contrib contrib-type="author"{corr}>')
        if auth.get('orcid'): L.append(f'          <contrib-id contrib-id-type="orcid">{xe(auth["orcid"])}</contrib-id>')
        L += ['          <name name-style="western">',
              f'            <surname>{xe(auth["surname"])}</surname>']
        if auth.get('given'): L.append(f'            <given-names>{xe(auth["given"])}</given-names>')
        L.append('          </name>')
        for an in auth.get('affiliationNums',[]):
            xid = nid(pfx,'x')
            L.append(f'          <xref id="{xid}" rid="aff{an}" ref-type="aff"><sup>{xe(an)}</sup></xref>')
        if auth.get('isCorresponding'):
            xid2 = nid(pfx,'x')
            L.append(f'          <xref id="{xid2}" rid="cor1" ref-type="corresp">*</xref>')
        L.append('        </contrib>')

    for num,txt in parsed.get('affiliations',{}).items():
        parts = [p.strip() for p in txt.split(',')]
        L.append(f'        <aff id="aff{num}">')
        L.append(f'          <label>{xe(num)}</label>')
        if len(parts) >= 2:
            L.append(f'          <institution content-type="dept">{xe(parts[0])}</institution>')
            mid = ', '.join(parts[1:-2]) if len(parts) > 3 else (parts[1] if len(parts) > 1 else '')
            if mid: L.append(f'          <institution>{xe(mid)}</institution>')
            if len(parts) >= 3: L.append(f'          <addr-line>{xe(parts[-2])}</addr-line>')
            L.append(f'          <country country="IN">{xe(parts[-1])}</country>')
        else: L.append(f'          <institution>{xe(txt)}</institution>')
        L.append('        </aff>')
    L.append('      </contrib-group>')
    L.append('')

    # author-notes
    corr_authors = [a for a in parsed.get('authors',[]) if a.get('isCorresponding')]
    L.append('      <author-notes>')
    if corr_authors:
        ca = corr_authors[0]
        L += ['        <corresp id="cor1">',
              f'          <bold>Corresponding Author:</bold> {xe(ca["given"])} {xe(ca["surname"])}',
              '        </corresp>']
    L += ['        <fn fn-type="coi-statement">','          <p>None declared.</p>','        </fn>',
          '        <fn fn-type="financial-disclosure">','          <p>None.</p>','        </fn>',
          '      </author-notes>','']

    # pub-date
    L += ['      <pub-date date-type="pub" publication-format="print">',
          f'        <day>{xe(jm.get("day",""))}</day>',
          f'        <month>{xe(jm.get("month",""))}</month>',
          f'        <year>{xe(jm.get("year","2025"))}</year>',
          '      </pub-date>']
    if jm.get('volume'): L.append(f'      <volume>{xe(jm["volume"])}</volume>')
    if jm.get('issue'):  L.append(f'      <issue>{xe(jm["issue"])}</issue>')
    if jm.get('fpage'):  L.append(f'      <fpage>{xe(jm["fpage"])}</fpage>')
    if jm.get('lpage'):  L.append(f'      <lpage>{xe(jm["lpage"])}</lpage>')

    # history
    if parsed.get('receivedDate') or parsed.get('acceptedDate'):
        L.append('      <history>')
        for dtype,dstr in [('received',parsed.get('receivedDate','')),('accepted',parsed.get('acceptedDate',''))]:
            if not dstr: continue
            d = dstr.replace('/','-').split('-')
            if len(d[0])==4: yr,mo,dy = d[0],d[1] if len(d)>1 else '',d[2] if len(d)>2 else ''
            else:            dy,mo,yr = d[0],d[1] if len(d)>1 else '',d[2] if len(d)>2 else ''
            L += [f'        <date date-type="{dtype}">',
                  f'          <day>{dy}</day>',f'          <month>{mo}</month>',
                  f'          <year>{yr}</year>','        </date>']
        L.append('      </history>')

    # permissions
    yr = jm.get('year','2025'); pub = jm.get('publisher','IP Innovative Publication')
    lic_url = jm.get('licenseUrl','https://creativecommons.org/licenses/by-nc/4.0/')
    L += ['      <permissions>',
          f'        <copyright-statement>© {yr} {xe(pub)}</copyright-statement>',
          f'        <copyright-year>{yr}</copyright-year>',
          f'        <copyright-holder>{xe(pub)}</copyright-holder>',
          f'        <license license-type="open-access" xlink:href="{lic_url}">',
          f'          <ali:license_ref>{lic_url}</ali:license_ref>',
          '          <license-p>This is an Open Access article distributed under the terms of the',
          f'          <ext-link ext-link-type="uri" xlink:href="{lic_url}" xlink:title="Creative Commons License">Creative Commons Attribution-NonCommercial 4.0 International License</ext-link>',
          '          which permits unrestricted non-commercial use, distribution, and reproduction',
          '          in any medium, provided the original work is properly cited.</license-p>',
          '        </license>','      </permissions>','']

    # abstract
    abstract = parsed.get('abstract',{})
    if abstract:
        abs_id = nid(pfx,'abstract')
        L += [f'      <abstract id="{abs_id}">','        <title>Abstract</title>']
        if len(abstract)==1 and 'text' in abstract:
            pid = nid(pfx,'p'); L.append(f'        <p id="{pid}">{xe(abstract["text"].strip())}</p>')
        else:
            for label,text in abstract.items():
                if label == 'text':
                    pid = nid(pfx,'p'); L.append(f'        <p id="{pid}">{xe(text.strip())}</p>')
                else:
                    sid=nid(pfx,'sec'); tid=nid(pfx,'title'); pid=nid(pfx,'p')
                    L += [f'        <sec id="{sid}">',f'          <title id="{tid}">{xe(label)}</title>',
                          f'          <p id="{pid}">{xe(text.strip())}</p>','        </sec>']
        L.append('      </abstract>')

    # keywords
    kws = parsed.get('keywords',[])
    if kws:
        kg_id = nid(pfx,'kwd-group')
        L += [f'      <kwd-group id="{kg_id}" kwd-group-type="author-generated">',
              '        <title>Keywords</title>']
        for kw in kws: L.append(f'        <kwd>{xe(kw)}</kwd>')
        L.append('      </kwd-group>')

    L += ['','    </article-meta>','  </front>','']

    # ── BODY ──
    L.append('  <body>')
    for sec in parsed.get('sections',[]):
        st = sec.get('sec_type')
        sec_attr = f' sec-type="{st}"' if st else ''
        L.append(f'    <sec{sec_attr}>')
        tid = nid(pfx,'title'); L.append(f'      <title id="{tid}">{xe(sec["title"])}</title>')

        for para in sec.get('paragraphs',[]):
            inline = para_to_inline(para, pfx, cite_ids, table_ids, fig_ids)
            if inline.strip():
                pid = nid(pfx,'p'); L.append(f'      <p id="{pid}">{inline}</p>')
            # Place tables inline after the para that first mentions them
            for tbl in parsed.get('tables',[]):
                if not tbl.get('placed') and f"Table {tbl['num']}" in para.text:
                    L += build_table_xml(tbl, pfx); tbl['placed'] = True
            # Place figures inline after the para that first mentions them
            for fig in parsed.get('figures',[]):
                if not fig.get('placed'):
                    ptxt = para.text
                    if re.search(rf'(Fig\.?\s*|Figure\s*){fig["num"]}\b', ptxt, re.I):
                        L += build_fig_xml(fig, pfx); fig['placed'] = True

        for sub in sec.get('subsections',[]):
            sst = sub.get('sec_type')
            sub_attr = f' sec-type="{sst}"' if sst else ''
            L.append(f'      <sec{sub_attr}>')
            stid = nid(pfx,'title'); L.append(f'        <title id="{stid}">{xe(sub["title"])}</title>')
            for para in sub.get('paragraphs',[]):
                inline = para_to_inline(para, pfx, cite_ids, table_ids, fig_ids)
                if inline.strip():
                    pid = nid(pfx,'p'); L.append(f'        <p id="{pid}">{inline}</p>')
                for tbl in parsed.get('tables',[]):
                    if not tbl.get('placed') and f"Table {tbl['num']}" in para.text:
                        L += build_table_xml(tbl, pfx); tbl['placed'] = True
                for fig in parsed.get('figures',[]):
                    if not fig.get('placed'):
                        if re.search(rf'(Fig\.?\s*|Figure\s*){fig["num"]}\b', para.text, re.I):
                            L += build_fig_xml(fig, pfx); fig['placed'] = True
            L.append('      </sec>')

        L.append('    </sec>'); L.append('')

    # Any unplaced tables at end of body
    for tbl in parsed.get('tables',[]):
        if not tbl.get('placed'):
            L += build_table_xml(tbl, pfx); tbl['placed'] = True
    # Any unplaced figures at end of body
    for fig in parsed.get('figures',[]):
        if not fig.get('placed'):
            L += build_fig_xml(fig, pfx); fig['placed'] = True

    L += ['  </body>','']

    # ── BACK ──
    L.append('  <back>')
    refs = parsed.get('references',[])
    if refs:
        L += ['    <ref-list>','      <title>References</title>']
        for ref in refs:
            p = ref.get('parsed',{}); cr = ref.get('crossref') or {}; pm = ref.get('pubmed') or {}
            # Priority: CrossRef > PubMed > local parse
            authors   = cr.get('authors') or pm.get('authors') or p.get('authors',[])
            title_r   = cr.get('title')   or pm.get('title')   or p.get('title','')
            journal_r = cr.get('journal') or pm.get('journal') or p.get('journal','')
            year_r    = cr.get('year')    or pm.get('year')    or p.get('year','')
            vol_r     = cr.get('volume')  or pm.get('volume')  or p.get('volume','')
            iss_r     = cr.get('issue')   or pm.get('issue')   or p.get('issue','')
            fp_r      = cr.get('fpage')   or pm.get('fpage')   or p.get('fpage','')
            lp_r      = cr.get('lpage')   or pm.get('lpage')   or p.get('lpage','')
            doi_raw   = ref.get('doi')    or cr.get('doi')     or pm.get('doi','')
            doi       = doi_raw if (doi_raw and '/' in doi_raw) else ''
            pmid      = pm.get('pmid','')
            pub_t     = p.get('pubType','journal')

            L += [f'      <ref id="{pfx}-B{ref["num"]}">',
                  f'        <label>{ref["num"]}.</label>',
                  f'        <element-citation publication-type="{pub_t}">']

            if authors:
                L.append('          <person-group person-group-type="author">')
                for a in authors:
                    L += ['            <name name-style="western">',
                          f'              <surname>{xe(a.get("surname",""))}</surname>']
                    if a.get('given'): L.append(f'              <given-names>{xe(a["given"])}</given-names>')
                    L.append('            </name>')
                if p.get('hasEtAl'): L.append('            <etal/>')
                L.append('          </person-group>')

            if title_r:   L.append(f'          <article-title>{xe(title_r)}</article-title>')
            if journal_r: L.append(f'          <source>{xe(journal_r)}</source>')
            if year_r:    L.append(f'          <year iso-8601-date="{year_r}">{xe(year_r)}</year>')
            if vol_r:     L.append(f'          <volume>{xe(vol_r)}</volume>')
            if iss_r:     L.append(f'          <issue>{xe(iss_r)}</issue>')
            if fp_r:      L.append(f'          <fpage>{xe(fp_r.strip())}</fpage>')
            if lp_r:      L.append(f'          <lpage>{xe(lp_r.strip())}</lpage>')
            if doi and '/' in doi:  L.append(f'          <pub-id pub-id-type="doi">{xe(doi)}</pub-id>')
            if pmid:      L.append(f'          <pub-id pub-id-type="pmid">{xe(pmid)}</pub-id>')
            if not authors and not title_r:
                L.append(f'          <!-- RAW: {xe(ref["raw"][:200])} -->')
            L += ['        </element-citation>','      </ref>']
        L.append('    </ref-list>')
    L += ['  </back>','','</article>']
    return '\n'.join(L)

def build_table_xml(tbl, pfx):
    L = []
    tw_id = f"tw-{tbl['num']}"
    L.append(f'      <table-wrap id="{tw_id}" position="anchor" orientation="portrait">')
    L.append(f'        <label>Table {tbl["num"]}</label>')
    cap = re.sub(r'^Table\s+\d+\s*[:\-]\s*','',tbl.get('caption',''),flags=re.I)
    if cap:
        cid=nid(pfx,'cap'); ctid=nid(pfx,'ctitle')
        L += [f'        <caption id="{cid}">',f'          <title id="{ctid}">{xe(cap)}</title>','        </caption>']
    tid = nid(pfx,'tbl')
    L.append(f'        <table id="{tid}" rules="all" frame="box">')
    rows=tbl.get('rows',[]); cws=tbl.get('colwidths',[])
    if cws:
        L.append('          <colgroup>')
        for w in cws: L.append(f'            <col width="{w}"/>')
        L.append('          </colgroup>')
    elif rows and rows[0]:
        n=len(rows[0]); w=round(100/n,2)
        L.append('          <colgroup>')
        for _ in rows[0]: L.append(f'            <col width="{w}"/>')
        L.append('          </colgroup>')
    if rows:
        L += [f'          <thead id="{nid(pfx,"thead")}">',f'            <tr id="{nid(pfx,"tr")}">']
        for cell in rows[0]:
            tc_id=nid(pfx,'th'); cs=f' colspan="{cell["colspan"]}"' if int(cell.get('colspan',1))>1 else ''
            ch=xe(cell['text'])
            L.append(f'              <th id="{tc_id}"{cs} align="left">')
            if ch.strip():
                L.append(f'                <p id="{nid(pfx,"p")}"><bold id="{nid(pfx,"b")}">{ch}</bold></p>')
            L.append('              </th>')
        L += ['            </tr>','          </thead>']
        if len(rows)>1:
            L.append(f'          <tbody id="{nid(pfx,"tbody")}">')
            for row in rows[1:]:
                L.append(f'            <tr id="{nid(pfx,"tr")}">')
                for cell in row:
                    td_id=nid(pfx,'td'); cs=f' colspan="{cell["colspan"]}"' if int(cell.get('colspan',1))>1 else ''
                    ct=xe(cell['text'])
                    L.append(f'              <td id="{td_id}"{cs} align="left">')
                    if ct.strip(): L.append(f'                <p id="{nid(pfx,"p")}">{ct}</p>')
                    L.append('              </td>')
                L.append('            </tr>')
            L.append('          </tbody>')
    L += ['        </table>','      </table-wrap>']
    return L

def build_fig_xml(fig, pfx):
    """Build JATS <fig> element"""
    L = []
    fig_id = f"fig-{fig['num']}"
    L.append(f'      <fig id="{fig_id}" position="anchor" orientation="portrait">')
    L.append(f'        <label>Figure {fig["num"]}</label>')
    cap = re.sub(r'^(Fig\.?\s*|Figure\s*)\d+\s*[:\-\.\s]*', '', fig.get('caption',''), flags=re.I).strip()
    if cap:
        cid = nid(pfx, 'cap'); ctid = nid(pfx, 'ctitle')
        L += [f'        <caption id="{cid}">',
              f'          <title id="{ctid}">{xe(cap)}</title>',
              '        </caption>']
    # Graphic placeholder (actual image not embedded — note added)
    L.append(f'        <graphic xlink:href="fig{fig['num']}" mimetype="image" mime-subtype="jpeg"/>')
    L.append('      </fig>')
    return L

def post_process(xml):
    xml = re.sub(r'<p([^>]*)>[\s\u00a0]*(?:&#160;|&#xA0;)?[\s\u00a0]*</p>','',xml)
    xml = re.sub(r'<bold([^>]*)>\s*</bold>','',xml)
    xml = re.sub(r'<italic([^>]*)>\s*</italic>','',xml)
    return '\n'.join(l for l in xml.split('\n') if l.strip())

# ── CLI ──────────────────────────────────────────────────────
if __name__=='__main__':
    import argparse
    ap=argparse.ArgumentParser(description='DOCX → JATS XML v4')
    ap.add_argument('input'); ap.add_argument('-o','--output')
    ap.add_argument('--journal',default=''); ap.add_argument('--issn-print',default='')
    ap.add_argument('--issn-elec',default=''); ap.add_argument('--publisher',default='IP Innovative Publication')
    ap.add_argument('--journal-url',default=''); ap.add_argument('--doi',default='')
    ap.add_argument('--volume',default=''); ap.add_argument('--issue',default='')
    ap.add_argument('--year',default='2025'); ap.add_argument('--month',default='')
    ap.add_argument('--day',default=''); ap.add_argument('--fpage',default=''); ap.add_argument('--lpage',default='')
    ap.add_argument('--type',default='research-article',dest='article_type')
    ap.add_argument('--license',default='cc-by-nc-4.0'); ap.add_argument('--no-crossref',action='store_true')
    args=ap.parse_args()

    TYPE_LABELS={'research-article':'Original Research Article','review-article':'Review Article',
        'case-report':'Case Report','letter':'Letter to Editor','editorial':'Editorial',
        'brief-report':'Brief Report','systematic-review':'Systematic Review'}
    LIC={'cc-by-4.0':'https://creativecommons.org/licenses/by/4.0/',
         'cc-by-nc-4.0':'https://creativecommons.org/licenses/by-nc/4.0/',
         'cc-by-nc-nd-4.0':'https://creativecommons.org/licenses/by-nc-nd/4.0/'}

    jm={'name':args.journal,'issnPrint':args.issn_print,'issnElec':args.issn_elec,
        'publisher':args.publisher,'journalUrl':args.journal_url,'doi':args.doi,
        'volume':args.volume,'issue':args.issue,'year':args.year,'month':args.month,
        'day':args.day,'fpage':args.fpage,'lpage':args.lpage,
        'articleType':args.article_type,'articleTypeLabel':TYPE_LABELS.get(args.article_type,args.article_type),
        'licenseUrl':LIC.get(args.license,LIC['cc-by-nc-4.0'])}

    print(f"📄 Parsing: {args.input}",file=sys.stderr)
    parsed=parse_docx(args.input,use_crossref=not args.no_crossref)
    xml=build_xml(parsed,jm); xml=post_process(xml)
    out=args.output or args.input.replace('.docx','-jats-v4.xml')
    with open(out,'w',encoding='utf-8') as f: f.write(xml)
    print(f"\n✅ {out} ({len(xml)/1024:.1f} KB | {xml.count(chr(10))} lines)",file=sys.stderr)
